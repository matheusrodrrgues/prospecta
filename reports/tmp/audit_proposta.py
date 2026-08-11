from pathlib import Path
import sys
from zipfile import ZipFile

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
DOCX = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "reports" / "output" / "Proposta_Prospecta_UNet_PSPNet_Fosfato_FINAL.docx"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


doc = Document(DOCX)
with ZipFile(DOCX) as archive:
    document_xml = etree.fromstring(archive.read("word/document.xml"))
    styles_xml = etree.fromstring(archive.read("word/styles.xml"))

page_breaks = document_xml.xpath("count(.//w:br[@w:type='page'])", namespaces=NS)
title_borders = document_xml.xpath("count(.//w:body/w:p[1]/w:pPr/w:pBdr)", namespaces=NS)
title_underlines = document_xml.xpath("count(.//w:body/w:p[1]//w:u)", namespaces=NS)

section = doc.sections[0]
assert int(page_breaks) == 2, f"Esperados 2 page breaks; encontrados {page_breaks}"
assert int(title_borders) == 0, "Título contém borda"
assert int(title_underlines) == 0, "Título contém sublinhado"
assert len(doc.tables) == 1, f"Esperada 1 tabela; encontradas {len(doc.tables)}"
assert len(doc.tables[0].rows) == 7, "Cronograma deve conter cabeçalho + 6 etapas"
assert len(doc.tables[0].columns) == 3, "Cronograma deve conter 3 colunas"

grid_widths = [int(n.get(f"{{{NS['w']}}}w")) for n in document_xml.xpath(".//w:tbl[1]/w:tblGrid/w:gridCol", namespaces=NS)]
assert grid_widths == [4320, 1440, 3600], grid_widths
assert sum(grid_widths) == 9360

all_text = "\n".join(p.text for p in doc.paragraphs)
assert "U-Net" in all_text and "PSPNet" in all_text
assert "[Nome]" not in all_text and "TODO" not in all_text
assert "PROSPECTA 4.0" in all_text

assert section.page_width.inches == 8.5
assert section.page_height.inches == 11.0
for margin in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
    assert margin.inches == 1.0

normal = doc.styles["Normal"]
assert normal.font.name == "Arial"
assert round(normal.font.size.pt, 1) == 11.0

words = len(all_text.split())
print({
    "arquivo": str(DOCX),
    "quebras_de_pagina": int(page_breaks),
    "paginas_estruturais": int(page_breaks) + 1,
    "palavras": words,
    "tabela": {"linhas": 7, "colunas": 3, "larguras_dxa": grid_widths},
    "titulo_sem_borda": True,
    "estilo_normal": "Arial 11 pt, 1.15",
})
