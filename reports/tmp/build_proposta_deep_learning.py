from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "reports" / "output" / "Proposta_Prospecta_UNet_PSPNet_Fosfato_REVISAO_ASTER_SEM_MAPA.docx"


def set_run_font(run, size=11, bold=False, italic=False, color="000000"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text, keep_with_next=False):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.add_run(text)
    return paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    tokens = {
        1: (20, 20, 6, "000000"),
        2: (16, 18, 6, "000000"),
        3: (14, 16, 4, "434343"),
    }
    for level, (size, before, after, color) in tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    # Página 1 — contextualização e objetivos.
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run("Avaliação comparativa das arquiteturas U-Net e PSPNet na prospecção de fosfato na Bacia de Irecê")
    set_run_font(title_run, size=26, bold=False)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.keep_with_next = True
    subtitle_run = subtitle.add_run("Proposta de plano de trabalho — PROSPECTA 4.0 | duração prevista: 12 meses")
    set_run_font(subtitle_run, size=11, color="555555")

    add_heading(doc, "1. Contextualização e justificativa", 1)
    add_body(doc, "O fosfato constitui matéria-prima essencial para a produção de fertilizantes e para a segurança alimentar. Na Bahia, a Bacia de Irecê reúne sequências carbonáticas do Grupo Una e registros históricos de fosforito, incluindo o Prospecto Central Irecê-Lapão. A extensão regional dessas unidades e a distribuição irregular de informações de campo tornam a delimitação de áreas favoráveis um problema que demanda a integração de dados geológicos, geoquímicos, topográficos e de sensoriamento remoto.")
    add_body(doc, "O ciclo anterior do PROSPECTA 4.0 estruturou uma plataforma web e geoespacial capaz de organizar ocorrências minerais, processar imagens orbitais no Google Earth Engine (GEE), armazenar produtos raster e disponibilizá-los em um mapa interativo. Também definiu as etapas necessárias para a modelagem de favorabilidade do fosfato com algoritmos tradicionais de Machine Learning. O avanço proposto neste plano é utilizar essa infraestrutura e o mosaico histórico ASTER de 2000–2008 já preparado como asset no GEE para investigar modelos de Deep Learning que aprendam padrões espaciais diretamente de conjuntos multicanais de evidências geocientíficas.")
    add_body(doc, "O uso de imagens ASTER é particularmente relevante porque o sensor combina bandas no visível e infravermelho próximo (VNIR), no infravermelho de ondas curtas (SWIR) e no infravermelho termal (TIR). As seis bandas SWIR, com resolução espacial de 30 m, são sensíveis a diferenças de composição mineralógica e podem contribuir para reconhecer litologias, minerais de alteração e contextos associados à mineralização. O intervalo histórico até 2008 possui importância adicional: após 23 de abril daquele ano, o subsistema SWIR apresentou aquecimento anômalo, saturação e listramento, tornando os dados posteriores inadequados. Assim, o mosaico pré-anomalia reúne a faixa temporal mais consistente para explorar conjuntamente as bandas espectrais de maior interesse geológico.")
    add_body(doc, "A U-Net combina um caminho de contração, responsável pela extração de contexto, com um caminho de expansão e conexões de atalho que favorecem a localização precisa dos alvos. A PSPNet, por sua vez, utiliza agrupamento piramidal para incorporar informações de contexto em diferentes escalas. Embora ambas realizem predição por pixel, seus mecanismos de representação espacial são distintos, o que justifica uma avaliação controlada no mapeamento de favorabilidade mineral. Estudos anteriores já demonstraram a viabilidade de redes profundas e da U-Net na modelagem de prospectividade; contudo, a aplicação comparativa dessas duas arquiteturas ao fosfato da Bacia de Irecê permanece uma questão a ser investigada.")

    add_heading(doc, "2. Objetivos", 1)
    add_heading(doc, "2.1 Objetivo geral", 2)
    add_body(doc, "Desenvolver e avaliar uma metodologia baseada nas arquiteturas U-Net e PSPNet para produzir mapas de favorabilidade à ocorrência de fosfato na Bacia de Irecê, comparando o desempenho preditivo, a coerência espacial e o custo computacional dos modelos e integrando os resultados à plataforma PROSPECTA 4.0.")
    add_heading(doc, "2.2 Objetivos específicos", 2)
    add_body(doc, "Organizar e harmonizar as camadas geoespaciais de entrada; construir amostras e máscaras de referência a partir de dados validados; implementar uma rotina reproduzível de treinamento e inferência; treinar U-Net e PSPNet sob condições comparáveis; avaliar os resultados com métricas de segmentação e critérios espaciais; produzir mapas de favorabilidade e incerteza; e disponibilizar as camadas resultantes no dashboard do PROSPECTA 4.0.")

    doc.add_page_break()

    # Página 2 — metodologia.
    add_heading(doc, "3. Metodologia", 1)
    add_heading(doc, "3.1 Área de estudo e base de evidências", 2)
    add_body(doc, "O estudo será conduzido na Bacia de Irecê, no centro-norte da Bahia, em área que abrange ocorrências e unidades geológicas relacionadas aos registros de fosfato do Prospecto Central Irecê–Lapão. A delimitação definitiva deverá acompanhar o polígono adotado no GEE e a cobertura das amostras validadas. O mapa de localização será produzido no próprio GEE, apresentando o limite da área de estudo sobre imagem de referência, a posição no estado da Bahia e, quando autorizada, a distribuição das ocorrências utilizadas. A redação e a delimitação cartográfica serão revisadas em conjunto com a equipe de geociências.")
    add_heading(doc, "3.2 Mosaico ASTER e preparação das evidências", 2)
    add_body(doc, "Será utilizado o asset do mosaico ASTER correspondente ao período de 2000 a abril de 2008, disponibilizado no ambiente GEE do projeto. O produto consolida cenas do período anterior à falha do subsistema SWIR, reduz lacunas espaciais e oferece uma base espectral contínua para a área de estudo. Antes da modelagem serão verificados a procedência do asset, as datas e cenas de origem, as bandas efetivamente presentes, o método de composição, a máscara de nuvens, a projeção e a resolução. A base multicanal deverá combinar bandas ASTER VNIR e SWIR válidas, composições e razões espectrais geologicamente justificadas, modelo digital de elevação e atributos do relevo. Sempre que disponíveis, serão adicionadas camadas de litologia, contatos, estruturas, distância a feições geológicas, ocorrências minerais e resultados geoquímicos de P₂O₅. Todas as fontes serão harmonizadas em uma grade comum, evitando atribuir precisão artificial às evidências menos detalhadas.")

    add_heading(doc, "3.3 Preparação das amostras", 2)
    add_body(doc, "As ocorrências e amostras confirmadas formarão a referência positiva. A representação das classes será definida com apoio da equipe de geociências, utilizando pontos, áreas de influência ou polígonos quando houver base técnica para isso. As amostras de ausência ou de fundo serão selecionadas de modo estratificado, afastadas das ocorrências conhecidas e distribuídas entre diferentes contextos geológicos. A base será dividida em blocos espaciais independentes de treinamento, validação e teste, reduzindo o vazamento causado pela proximidade entre pixels. Recortes multicanais serão extraídos como entrada, com aumentos de dados por rotações e espelhamentos aplicados somente ao conjunto de treinamento.")

    add_heading(doc, "3.4 Modelos e treinamento", 2)
    add_body(doc, "Serão implementadas duas redes de segmentação semântica. A U-Net atuará como referência encoder-decoder, preservando detalhes locais pelas conexões entre níveis equivalentes. A PSPNet será configurada com seu módulo de agrupamento piramidal, destinado a combinar contexto local e regional. Os modelos receberão as mesmas camadas, divisões espaciais e orçamento de treinamento. O otimizador Adam e uma taxa inicial próxima de 3 × 10⁻⁴ serão avaliados como ponto de partida, conforme a síntese bibliográfica da planilha do projeto. Funções de perda ponderadas, Focal Loss e combinações com Dice serão testadas para tratar o desbalanceamento; a escolha final será feita exclusivamente com o conjunto de validação.")

    add_heading(doc, "3.5 Avaliação comparativa", 2)
    add_body(doc, "O desempenho no teste espacial será medido por precisão, revocação, F1/Dice, interseção sobre união (IoU), área sob a curva ROC e, prioritariamente em cenário desbalanceado, área sob a curva precisão-revocação. Também serão registrados tempo de treinamento, memória utilizada e número de parâmetros. A interpretação geográfica considerará a proporção de ocorrências conhecidas recuperadas nas classes de maior favorabilidade, a área ocupada por essas classes e a continuidade dos padrões produzidos. Os mapas não serão apresentados como confirmação de jazidas, mas como superfícies probabilísticas de apoio à seleção de alvos para verificação geológica.")

    add_heading(doc, "3.6 Integração ao PROSPECTA 4.0", 2)
    add_body(doc, "As previsões serão exportadas em formato raster georreferenciado, acompanhadas por metadados de versão, data, camadas de entrada, hiperparâmetros e métricas. O melhor modelo e a diferença entre as previsões das duas arquiteturas serão publicados como camadas no dashboard, permitindo controlar opacidade, consultar valores por local e sobrepor ocorrências. A integração manterá separadas a visualização exploratória, a saída científica validada e os registros comunitários já suportados pela plataforma.")

    doc.add_page_break()

    # Página 3 — produtos, cronograma e referências.
    add_heading(doc, "4. Resultados esperados", 1)
    add_body(doc, "Espera-se obter uma base multicanal documentada para a Bacia de Irecê, uma pipeline reproduzível de treinamento e inferência, dois modelos avaliados sob protocolo espacial comum e mapas de favorabilidade acompanhados de métricas e limitações. A comparação deverá indicar em quais condições a preservação de detalhes da U-Net ou a agregação multiescala da PSPNet é mais adequada ao problema. Como produto tecnológico, o PROSPECTA 4.0 receberá camadas preditivas reais e versionadas, substituindo superfícies meramente demonstrativas. Como produto científico, serão produzidos relatório final, apresentação no SEMIC e manuscrito ou resumo técnico sobre a avaliação comparativa.")

    add_heading(doc, "5. Cronograma de execução", 1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [4320, 1440, 3600])
    set_table_borders(table)
    headers = ("Etapa", "Meses", "Produto principal")
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_tr_pr.append(header_flag)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=10, bold=True)

    schedule = [
        ("Revisão bibliográfica e definição do recorte", "1–2", "Protocolo do experimento"),
        ("Aquisição, integração e auditoria dos dados", "2–4", "Base geoespacial harmonizada"),
        ("Construção das máscaras e divisão espacial", "4–5", "Amostras de treino, validação e teste"),
        ("Implementação e treinamento dos modelos", "5–8", "U-Net e PSPNet treinadas"),
        ("Avaliação, análise espacial e ajustes", "8–10", "Métricas e mapas comparativos"),
        ("Integração, documentação e divulgação", "10–12", "Dashboard, relatório e SEMIC"),
    ]
    for stage, months, product in schedule:
        cells = table.add_row().cells
        for idx, text in enumerate((stage, months, product)):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=9.5)

    add_heading(doc, "6. Referências", 1)
    references = [
        "BONFIM, L. F. C. et al. Projeto Bacia de Irecê: relatório final. Salvador: CPRM/SGM, 1985.",
        "LIU, C. et al. A deep-learning-based mineral prospectivity modeling framework and workflow in prediction of porphyry–epithermal mineralization in the Duolong Ore District, Tibet. Ore Geology Reviews, v. 157, 105419, 2023. DOI: 10.1016/j.oregeorev.2023.105419.",
        "RONNEBERGER, O.; FISCHER, P.; BROX, T. U-Net: convolutional networks for biomedical image segmentation. In: MICCAI 2015. Cham: Springer, 2015. p. 234–241. DOI: 10.1007/978-3-319-24574-4_28.",
        "SERVIÇO GEOLÓGICO DO BRASIL. Prospecto Central Irecê-Lapão: relatório 02 — fosfato em Irecê. Brasília: SGB/CPRM.",
        "JET PROPULSION LABORATORY. ASTER instrument characteristics. Pasadena: NASA/JPL. Disponível em: https://asterweb.jpl.nasa.gov/content/01_mission/03_instrument/01_Characteristics/default.htm.",
        "JET PROPULSION LABORATORY. SWIR — ASTER User Advisory. Pasadena: NASA/JPL, 2008–2009. Disponível em: https://asterweb.jpl.nasa.gov/swir-alert.asp.",
        "GOOGLE. ASTER L1T Radiance: Earth Engine Data Catalog. Disponível em: https://developers.google.com/earth-engine/datasets/catalog/ASTER_AST_L1T_003.",
        "ZHAO, H. et al. Pyramid Scene Parsing Network. In: IEEE Conference on Computer Vision and Pattern Recognition, 2017. p. 2881–2890.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(ref)
        set_run_font(run, size=9.5)

    # Metadados neutros para compartilhamento.
    props = doc.core_properties
    props.title = "Avaliação comparativa das arquiteturas U-Net e PSPNet na prospecção de fosfato na Bacia de Irecê"
    props.subject = "Proposta de plano de trabalho do PROSPECTA 4.0"
    props.author = "PROSPECTA 4.0"
    props.keywords = "fosfato; Bacia de Irecê; Deep Learning; U-Net; PSPNet; prospecção mineral"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
