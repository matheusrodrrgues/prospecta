from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


ROOT = Path(__file__).resolve().parents[2]
REFERENCE = ROOT / "reports" / "input" / "modelo_semic2025.docx"
OUTPUT = ROOT / "reports" / "output" / "Relatorio_Prospecta_Fosfato_Irece_Machine_Learning.docx"


def set_font(run, name="Times New Roman", size=12, bold=None, italic=None, underline=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def remove_body_after(document, keep_paragraph_count=13):
    body = document._element.body
    kept = 0
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p") and kept < keep_paragraph_count:
            kept += 1
            continue
        body.remove(child)


def style_paragraph(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True,
                    before=0, after=0, keep_with_next=False):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.15
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(1.25) if first_line else Cm(0)
    fmt.keep_with_next = keep_with_next


def add_text_paragraph(document, text, *, first_line=True, after=0):
    p = document.add_paragraph()
    style_paragraph(p, first_line=first_line, after=after)
    set_font(p.add_run(text))
    return p


def add_section(document, title):
    p = document.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=5, after=0, keep_with_next=True)
    set_font(p.add_run(title), bold=True)
    return p


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_status_table(document):
    caption = document.add_paragraph()
    style_paragraph(caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=4, after=3, keep_with_next=True)
    set_font(caption.add_run("Tabela 1. Síntese das etapas implementadas e do estágio atual."), size=10)

    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = 1
    widths = [Cm(3.6), Cm(7.4), Cm(3.2)]
    headers = ["Etapa", "Produto/funcionalidade", "Situação"]
    for i, (cell, width, label) in enumerate(zip(table.rows[0].cells, widths, headers)):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "D9EAD3")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), size=9, bold=True)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

    rows = [
        ("Extração", "Mosaicos Landsat semestrais e exportação GeoTIFF otimizada para nuvem.", "Implementada"),
        ("Pré-processamento", "Máscara QA_PIXEL, reflectância de superfície, NDVI e qualityMosaic.", "Implementado"),
        ("Integração", "GCS/COG, metadados no PostGIS, API e visualização no GeoLab.", "Implementada"),
        ("Machine Learning", "Treinamento, validação e mapa probabilístico específico para fosfato.", "Próxima etapa"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, (cell, width, value) in enumerate(zip(cells, widths, values)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(value), size=9)
    p = document.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, after=0)
    set_font(p.add_run("Fonte: elaboração própria a partir do repositório Prospecta 4.0 (2026)."), size=9)


def add_reference(document, text):
    p = document.add_paragraph()
    style_paragraph(p, first_line=False, after=4)
    set_font(p.add_run(text), size=10)


OUTPUT.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(REFERENCE, OUTPUT)
doc = Document(OUTPUT)

# Preserve the institutional header and replace only the documented slots.
slots = doc.paragraphs
clear_paragraph(slots[5])
style_paragraph(slots[5], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, after=0)
set_font(slots[5].add_run("Modelagem do Potencial Mineral para Fosfato na Bacia de Irecê usando Machine Learning"), bold=True)

clear_paragraph(slots[7])
style_paragraph(slots[7], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
set_font(slots[7].add_run("Adriele Ramos Santana¹"), bold=True, underline=True)
set_font(slots[7].add_run("; Washington de Jesus Sant'Anna da Franca Rocha²"), bold=True)

clear_paragraph(slots[8])
style_paragraph(slots[8], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
set_font(slots[8].add_run("1. Adriele Ramos Santana, PIBITI/CNPq, Agronomia, e-mail: adrieleagronomia@gmail.com"), size=9)

clear_paragraph(slots[9])
style_paragraph(slots[9], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
set_font(slots[9].add_run("2. Washington de Jesus Sant'Anna da Franca Rocha, Departamento de Ciências Exatas, e-mail: wrocha@uefs.br"), size=9)

clear_paragraph(slots[11])
style_paragraph(slots[11], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
set_font(slots[11].add_run("PALAVRAS-CHAVE: "), bold=True)
set_font(slots[11].add_run("fosfato; sensoriamento remoto; Bacia de Irecê; aprendizado de máquina; Google Earth Engine"))

remove_body_after(doc, 13)

add_section(doc, "INTRODUÇÃO")
add_text_paragraph(doc, "O fosfato é uma matéria-prima estratégica para a produção de fertilizantes e, por consequência, para a segurança alimentar. Na Bahia, a Bacia de Irecê reúne sequências carbonáticas do Grupo Una e registros históricos de fosforito, incluindo o Prospecto Central Irecê-Lapão, descoberto durante levantamentos geológicos conduzidos na década de 1980. A extensão regional das unidades, a diversidade de controles geológicos e a distribuição irregular dos dados de campo tornam a avaliação de favorabilidade mineral um problema adequado à integração de sensoriamento remoto, sistemas de informação geográfica e métodos computacionais.")
add_text_paragraph(doc, "Nesse contexto, o Prospecta 4.0 vem sendo estruturado como uma plataforma web e geoespacial para organizar dados minerais, processar séries de imagens orbitais e disponibilizar resultados em um ambiente interativo. A hipótese de trabalho é que variáveis espectrais, temporais, geológicas e espaciais possam, após padronização e associação a ocorrências conhecidas de fosfato, alimentar modelos supervisionados capazes de estimar a favorabilidade mineral. O objetivo deste relatório é registrar as atividades realizadas até o momento: extração e pré-processamento de imagens Landsat, produção de mosaicos multitemporais, organização da infraestrutura de dados e desenvolvimento das ferramentas de exploração e análise, além de delimitar as etapas ainda necessárias para o treinamento e a validação do modelo de Machine Learning.")

add_section(doc, "METODOLOGIA")
add_text_paragraph(doc, "A infraestrutura foi organizada em três camadas: processamento geoespacial, persistência de dados e aplicação web. A área operacional inicial corresponde a um retângulo envolvendo o estado da Bahia (46,5° O a 37,0° O; 18,5° S a 8,5° S), enquanto a Bacia de Irecê é tratada no painel como recorte prioritário, com centro aproximado em 41,85° O e 11,30° S. Esse desenho permite produzir uma base estadual comum e, posteriormente, recortar as variáveis para a área efetiva de modelagem.")
add_text_paragraph(doc, "A extração orbital foi implementada no Google Earth Engine (GEE) com imagens Landsat Collection 2 Level-2. Foram utilizadas as coleções Landsat 5 TM para 2000-2011, Landsat 7 ETM+ para 2012, Landsat 8 OLI a partir de 2013 e a combinação Landsat 8/9 nos períodos recentes. Os pixels afetados por nuvem, sombra de nuvem e dilatação/cirrus são removidos a partir dos bits 3, 4 e 1 da banda QA_PIXEL. Em seguida, aplica-se o fator de escala 0,0000275 e o deslocamento -0,2 à reflectância de superfície, conforme a especificação da Collection 2. O índice de vegetação por diferença normalizada (NDVI) é calculado com as bandas vermelha e infravermelha próxima e utilizado no método qualityMosaic, que seleciona, para cada pixel, a observação de maior NDVI entre as cenas disponíveis.")
add_text_paragraph(doc, "Foram preparados mosaicos semestrais, com resolução espacial de 30 m, sistema de referência EPSG:4326 e saída GeoTIFF otimizada para nuvem (COG). O script histórico cobre um mosaico combinado de 2000-2008 e períodos semestrais entre 2009 e 2020, exceto 2012_1, marcado como indisponível por insuficiência de dados. A pipeline automatizada também identifica o último semestre completo, consulta se o produto já está pronto, calcula número de cenas e cobertura média de nuvens, exporta o COG para o Google Cloud Storage e registra metadados no Supabase/PostgreSQL com extensão PostGIS. O processamento é executável como Cloud Run Job, com controle de estado, tempo limite, mensagens de erro e modo incremental para evitar reprocessamento.")
add_text_paragraph(doc, "Na aplicação, os dados são disponibilizados por API e visualizados com MapLibre GL JS. O GeoLab permite selecionar região, categoria mineral, sensor, período, composição e mapa-base; controlar a opacidade; consultar ocorrências; importar até 500 pontos em CSV ou GeoJSON; filtrar valores por limiar; compartilhar a vista e exportar as feições visíveis. A arquitetura aceita tiles derivados dos COGs por meio de TiTiler e registra contribuições externas com separação entre dados institucionais e comunitários.")

add_section(doc, "RESULTADOS E/OU DISCUSSÃO")
add_text_paragraph(doc, "O principal resultado alcançado foi a consolidação de uma cadeia reproduzível entre o catálogo orbital e o ambiente de exploração. O pré-processamento uniformiza sensores de diferentes gerações, reduz a interferência de nuvens e produz mosaicos comparáveis em uma grade de 30 m. A série demonstrativa cadastrada no sistema contém 24 períodos, incluindo o agregado 2000-2008 e os semestres de 2009 a 2020 disponíveis. Cada período possui campos de qualidade, cobertura de nuvens, número de cenas, endereço do COG ou dos tiles e situação de processamento, permitindo rastrear a proveniência do produto.")
add_status_table(doc)
add_text_paragraph(doc, "Para o tema fosfato, a Bacia de Irecê já aparece como alvo estratégico em estudo e pode ser isolada no painel. A plataforma também oferece mecanismos úteis à etapa de preparação das amostras, como importação de pontos, sobreposição com ocorrências e exportação em GeoJSON. A superfície analítica atualmente exibida no mapa, entretanto, é uma grade exploratória sintética gerada no navegador a partir de um limiar; ela serve para testar interação, simbologia e fluxo de trabalho, mas não representa uma probabilidade geológica calculada por modelo treinado. Da mesma forma, os valores de qualidade do conjunto de demonstração não devem ser interpretados como métricas de acurácia mineral.")
add_text_paragraph(doc, "A etapa de Machine Learning permanece em desenvolvimento. Para convertê-la em modelagem científica será necessário reunir rótulos positivos confiáveis (ocorrências, amostras e teores de P2O5), definir amostras de ausência ou background, recortar e harmonizar os preditores e controlar o desbalanceamento de classes. Além das bandas e razões espectrais, o conjunto de variáveis deverá incorporar litologia, estruturas, distância a contatos, relevo e dados geoquímicos quando disponíveis. Recomenda-se comparar algoritmos interpretáveis, como regressão logística e Random Forest, com métodos de gradient boosting, usando validação espacial em blocos para reduzir a autocorrelação entre treino e teste. O desempenho deve ser avaliado por precisão, revocação, F1, área sob a curva ROC e área sob a curva precisão-revocação, acompanhado de análise de importância das variáveis e verificação geológica dos alvos.")
add_text_paragraph(doc, "Assim, o estágio atual não corresponde ainda a um mapa final de potencial mineral, mas a uma base tecnológica preparada para recebê-lo. Essa distinção é importante: a automação existente resolve tarefas de aquisição, armazenamento, visualização e rastreabilidade, enquanto a validade preditiva dependerá da qualidade dos dados de treinamento, do desenho da validação e da interpretação geológica. O avanço mais imediato é substituir a superfície sintética por uma camada derivada de dados reais e versionada junto aos metadados do experimento.")

add_section(doc, "CONSIDERAÇÕES FINAIS")
add_text_paragraph(doc, "As atividades realizadas no Prospecta 4.0 estabeleceram os componentes essenciais para a modelagem do potencial de fosfato na Bacia de Irecê: extração multitemporal Landsat, máscara de qualidade, cálculo de NDVI, composição por qualityMosaic, exportação COG, armazenamento de metadados espaciais e visualização interativa. O sistema já permite organizar períodos, ocorrências e contribuições, além de apoiar a inspeção e o intercâmbio de dados. O próximo ciclo deverá concentrar-se na construção do banco de amostras geológicas, na engenharia de atributos, no treinamento comparativo de modelos e na validação espacial. Somente após essas etapas será possível apresentar um mapa probabilístico de favorabilidade mineral com métricas de desempenho e incerteza adequadamente documentadas.")

add_section(doc, "REFERÊNCIAS")
add_reference(doc, "BONFIM, L. F. C. et al. 1985. Projeto Bacia de Irecê: relatório final. Salvador: CPRM/SGM.")
add_reference(doc, "BREIMAN, L. 2001. Random forests. Machine Learning, v. 45, p. 5-32.")
add_reference(doc, "GORELICK, N. et al. 2017. Google Earth Engine: planetary-scale geospatial analysis for everyone. Remote Sensing of Environment, v. 202, p. 18-27.")
add_reference(doc, "SERVIÇO GEOLÓGICO DO BRASIL (SGB/CPRM). Prospecto Central Irecê-Lapão: relatório 02 - fosfato em Irecê. Brasília: SGB/CPRM.")
add_reference(doc, "UNITED STATES GEOLOGICAL SURVEY (USGS). Landsat Collection 2 Level-2 Science Products. Reston: USGS.")

# Keep the original A4 geometry and force Word to refresh pagination fields if any.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.title = "Modelagem do Potencial Mineral para Fosfato na Bacia de Irecê usando Machine Learning"
doc.core_properties.subject = "Relatório de atividades do Prospecta 4.0"

# Accessibility metadata for the two preserved institutional marks.
for index, doc_pr in enumerate(doc._element.xpath(".//wp:docPr"), start=1):
    doc_pr.set("descr", "Brasão da UEFS" if index == 1 else "Marca da Pró-Reitoria de Pesquisa e Pós-Graduação da UEFS")
doc.save(OUTPUT)
print(OUTPUT)
