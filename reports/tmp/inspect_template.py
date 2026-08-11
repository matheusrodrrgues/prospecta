from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(__file__).parents[1] / "input" / "modelo_semic2025.docx"
doc = Document(path)

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.replace("\n", "\\n")
    drawing_count = len(paragraph._p.xpath(".//w:drawing"))
    p = paragraph.paragraph_format
    print(
        f"{i:03d} style={paragraph.style.name!r} align={paragraph.alignment} "
        f"drawings={drawing_count} before={p.space_before} after={p.space_after} "
        f"line={p.line_spacing} text={text!r}"
    )
    for j, run in enumerate(paragraph.runs):
        if run.text or run._r.xpath(".//w:drawing"):
            fonts = run._r.rPr.rFonts if run._r.rPr is not None else None
            print(
                f"    run {j}: text={run.text!r} bold={run.bold} italic={run.italic} "
                f"underline={run.underline} size={run.font.size} name={run.font.name!r} "
                f"ascii={fonts.get(qn('w:ascii')) if fonts is not None else None!r}"
            )

print(f"tables={len(doc.tables)} sections={len(doc.sections)}")
